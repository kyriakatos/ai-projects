#!/usr/bin/env python3
from __future__ import annotations

import argparse
import logging
import os
import re
import sys
from dataclasses import dataclass
from typing import List, Tuple

from docx import Document
from llama_cpp import Llama


# -------------------------
# Logging
# -------------------------
def setup_logging(level: str = "INFO") -> None:
    logging.basicConfig(
        level=getattr(logging, level.upper(), logging.INFO),
        format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
        handlers=[logging.StreamHandler(sys.stdout)],
    )


log = logging.getLogger("meltemi_summary_agent")


# -------------------------
# DOCX IO
# -------------------------
def read_docx_text(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    text = "\n".join(parts)
    text = re.sub(r"[ \t]+", " ", text).strip()
    return text


def write_docx(path: str, title: str, body: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    doc.save(path)


# -------------------------
# Utility
# -------------------------
def approx_word_count(text: str) -> int:
    # Greek/Latin words: rough tokenization by whitespace
    return len([w for w in re.split(r"\s+", text.strip()) if w])


def chunk_text(text: str, max_chars: int) -> List[str]:
    paras = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    buf = []
    size = 0
    for p in paras:
        if size + len(p) + 1 > max_chars and buf:
            chunks.append("\n".join(buf))
            buf = [p]
            size = len(p) + 1
        else:
            buf.append(p)
            size += len(p) + 1
    if buf:
        chunks.append("\n".join(buf))
    return chunks


# -------------------------
# LLM wrapper (Meltemi via llama.cpp)
# -------------------------
@dataclass
class LlamaParams:
    n_ctx: int = 8192
    n_threads: int = max(2, os.cpu_count() or 4)
    # Metal acceleration
    n_gpu_layers: int = -1
    main_gpu: int = 0
    n_batch: int = 1024
    f16_kv: bool = True
    temperature: float = 0.2
    top_p: float = 0.9


SYSTEM = (
    "Είσαι ένας αυστηρός, αξιόπιστος επιμελητής κειμένων στα Ελληνικά. "
    "Δεν επινοείς πληροφορίες. Προτιμάς σαφή δομή και πιστότητα."
)


def llm_complete(llm: Llama, prompt: str, max_tokens: int, temperature: float, top_p: float) -> str:
    out = llm(
        prompt,
        max_tokens=max_tokens,
        temperature=temperature,
        top_p=top_p,
        stop=["</s>"],
    )
    return (out["choices"][0]["text"] or "").strip()


# -------------------------
# Agent steps
# -------------------------
def agent_plan(llm: Llama, text: str, target_words: int) -> Tuple[int, int]:
    """
    Agent decides chunk size (chars) and how many revision loops to allow.
    Keeps it conservative and fast.
    """
    prompt = (
        f"{SYSTEM}\n\n"
        "Θα σου δώσω ΜΟΝΟ στατιστικά για ένα έγγραφο.\n"
        "Δώσε ένα σχέδιο εκτέλεσης με 2 τιμές:\n"
        "1) προτεινόμενο max_chars για chunking (αριθμός)\n"
        "2) μέγιστος αριθμός κύκλων αναθεώρησης (1-3)\n\n"
        f"Στατιστικά:\n"
        f"- characters: {len(text)}\n"
        f"- approx_words: {approx_word_count(text)}\n"
        f"- target_words_for_summary: {target_words}\n\n"
        "Απάντησε ΜΟΝΟ ως: max_chars=<int>; revisions=<int>"
    )
    resp = llm_complete(llm, prompt, max_tokens=80, temperature=0.0, top_p=1.0)
    m1 = re.search(r"max_chars\s*=\s*(\d+)", resp)
    m2 = re.search(r"revisions\s*=\s*(\d+)", resp)
    max_chars = int(m1.group(1)) if m1 else 9000
    revisions = int(m2.group(1)) if m2 else 2
    max_chars = max(3000, min(14000, max_chars))
    revisions = max(1, min(3, revisions))
    return max_chars, revisions


def summarize_chunks_to_bullets(llm: Llama, chunks: List[str]) -> str:
    bullet_blocks = []
    for i, ch in enumerate(chunks, start=1):
        log.info("Chunk %d/%d: summarizing to bullets", i, len(chunks))
        prompt = (
            f"{SYSTEM}\n\n"
            "Σύνοψισε το παρακάτω απόσπασμα στα Ελληνικά σε 8–12 κουκκίδες.\n"
            "- Κράτησε βασικά σημεία, συμπεράσματα, αριθμούς/ονόματα.\n"
            "- Αφαίρεσε επαναλήψεις.\n\n"
            f"Κείμενο:\n{ch}\n"
        )
        bullets = llm_complete(llm, prompt, max_tokens=450, temperature=0.2, top_p=0.9)
        bullet_blocks.append(bullets)
    return "\n\n".join(bullet_blocks)


def draft_one_page_summary(llm: Llama, bullets_joined: str, target_words: int) -> str:
    prompt = (
        f"{SYSTEM}\n\n"
        "Με βάση ΜΟΝΟ τις κουκκίδες που ακολουθούν, γράψε περίληψη ~1 σελίδας στα Ελληνικά.\n"
        f"Στόχος μήκους: περίπου {target_words} λέξεις.\n"
        "Μορφή:\n"
        "1) Εισαγωγή 2–3 προτάσεων.\n"
        "2) 5–8 βασικά σημεία σε κουκκίδες.\n"
        "3) Τελική παράγραφος συμπεράσματος.\n"
        "Κανόνες: μη προσθέτεις νέα στοιχεία.\n\n"
        f"Κουκκίδες:\n{bullets_joined}\n"
    )
    return llm_complete(llm, prompt, max_tokens=1100, temperature=0.2, top_p=0.9)


def critique_summary(llm: Llama, bullets_joined: str, summary: str, target_words: int) -> str:
    prompt = (
        f"{SYSTEM}\n\n"
        "Έλεγξε την περίληψη σε σχέση με τις κουκκίδες.\n"
        "Βγάλε ένα σύντομο report με:\n"
        "- Σφάλματα πιστότητας (αν η περίληψη λέει κάτι που δεν υπάρχει στις κουκκίδες)\n"
        "- Κύρια σημεία που λείπουν\n"
        "- Αν το μήκος είναι πολύ μικρό/μεγάλο (στόχος περίπου {target_words} λέξεις)\n"
        "Κράτησε το report κάτω από 180 λέξεις.\n\n"
        f"Κουκκίδες:\n{bullets_joined}\n\n"
        f"Περίληψη:\n{summary}\n"
    )
    return llm_complete(llm, prompt, max_tokens=280, temperature=0.0, top_p=1.0)


def revise_summary(llm: Llama, bullets_joined: str, summary: str, critique: str, target_words: int) -> str:
    prompt = (
        f"{SYSTEM}\n\n"
        "Αναθεώρησε την περίληψη σύμφωνα με το report.\n"
        f"Στόχος: ~{target_words} λέξεις, ίδια μορφή (εισαγωγή, κουκκίδες, συμπέρασμα).\n"
        "Κανόνες: μην επινοείς.\n\n"
        f"Report:\n{critique}\n\n"
        f"Κουκκίδες:\n{bullets_joined}\n\n"
        f"Τρέχουσα περίληψη:\n{summary}\n"
    )
    return llm_complete(llm, prompt, max_tokens=1100, temperature=0.2, top_p=0.9)


# -------------------------
# Main
# -------------------------
def main() -> None:
    parser = argparse.ArgumentParser(description="AI agent summarizer: Greek DOCX -> 1-page Greek summary (Meltemi + Metal).")
    parser.add_argument("--model", required=True, help="Path to Meltemi GGUF (e.g. meltemi-7b-instruct.Q5_K_M.gguf)")
    parser.add_argument("--input", required=True, help="Input .docx (Greek)")
    parser.add_argument("--output", required=True, help="Output summary .docx")
    parser.add_argument("--target-words", type=int, default=520, help="Approx words for 1 page (default ~520)")
    parser.add_argument("--log-level", default=os.getenv("LOG_LEVEL", "INFO"))
    args = parser.parse_args()

    setup_logging(args.log_level)

    text = read_docx_text(args.input)
    if not text:
        raise SystemExit("Το έγγραφο φαίνεται άδειο ή μη αναγνώσιμο.")

    log.info("Loading Meltemi model with Metal...")
    params = LlamaParams()
    llm = Llama(
        model_path=args.model,
        n_ctx=params.n_ctx,
        n_threads=params.n_threads,
        n_gpu_layers=params.n_gpu_layers,
        main_gpu=params.main_gpu,
        n_batch=params.n_batch,
        use_mmap=True,
        use_mlock=False,
        f16_kv=params.f16_kv,
        verbose=False,
    )

    # Agent planning
    max_chars, max_revisions = agent_plan(llm, text, args.target_words)
    log.info("Agent plan: max_chars=%d, max_revisions=%d", max_chars, max_revisions)

    # Chunk + bullet map
    chunks = chunk_text(text, max_chars=max_chars)
    log.info("Chunking: chars=%d -> chunks=%d", len(text), len(chunks))
    bullets = summarize_chunks_to_bullets(llm, chunks)

    # Draft + critique + revise loop
    summary = draft_one_page_summary(llm, bullets, args.target_words)
    for i in range(1, max_revisions + 1):
        wc = approx_word_count(summary)
        log.info("Draft word count ~%d (target ~%d). Critique pass %d/%d", wc, args.target_words, i, max_revisions)
        critique = critique_summary(llm, bullets, summary, args.target_words)

        # If critique says it's fine and length is close-ish, stop early
        close_enough = abs(wc - args.target_words) <= max(80, int(args.target_words * 0.15))
        looks_ok = "σφάλ" not in critique.lower() and "λάθος" not in critique.lower()
        if close_enough and looks_ok:
            log.info("Stopping early: length close and critique clean enough.")
            break

        summary = revise_summary(llm, bullets, summary, critique, args.target_words)

    write_docx(args.output, title="Περίληψη (1 σελίδα)", body=summary)
    log.info("Wrote summary DOCX: %s", args.output)


if __name__ == "__main__":
    main()