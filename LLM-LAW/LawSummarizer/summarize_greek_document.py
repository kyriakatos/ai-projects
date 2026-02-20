#!/usr/bin/env python3
"""
Summarize a Greek DOCX into ~1 page Greek summary using Meltemi (GGUF) via llama-cpp.

Outputs:
- summary.docx (Greek summary)
- summary.txt (optional)

Usage:
  python summarize_greek_docx_meltemi.py \
    --model /path/to/meltemi.gguf \
    --input /path/to/input.docx \
    --output /path/to/summary.docx
"""

from __future__ import annotations

import argparse
import logging
import os
import re
import sys
from dataclasses import dataclass
from typing import List

from docx import Document
from llama_cpp import Llama


# -------------------------
# Logging
# -------------------------
def setup_logging(level: str = "INFO") -> None:
    logging.basicConfig(
        level=getattr(logging, level.upper(), logging.INFO),
        format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
        handlers=[logging.StreamHandler(sys.stdout)],
    )


log = logging.getLogger("meltemi_docx_summarizer")


# -------------------------
# DOCX utilities
# -------------------------
def read_docx_text(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    text = "\n".join(parts)
    # normalize whitespace
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def write_docx(path: str, title: str, body: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n\n"):
        doc.add_paragraph(para.strip())
    doc.save(path)


# -------------------------
# Chunking (keeps prompts within context)
# -------------------------
def chunk_text(text: str, max_chars: int = 9000) -> List[str]:
    """
    Simple chunker by paragraphs aiming for max_chars per chunk.
    For Greek DOCX this is usually adequate.
    """
    paras = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    buf = []
    size = 0
    for p in paras:
        if size + len(p) + 1 > max_chars and buf:
            chunks.append("\n".join(buf))
            buf = [p]
            size = len(p) + 1
        else:
            buf.append(p)
            size += len(p) + 1
    if buf:
        chunks.append("\n".join(buf))
    return chunks


# -------------------------
# LLM prompting
# -------------------------
SYSTEM_STYLE = (
    "Είσαι ένας εξαιρετικός επιμελητής κειμένων στα Ελληνικά. "
    "Παράγεις σαφείς, δομημένες περιλήψεις με πιστότητα στο αρχικό κείμενο. "
    "Δεν επινοείς πληροφορίες."
)

def make_chunk_prompt(chunk: str) -> str:
    return (
        f"{SYSTEM_STYLE}\n\n"
        "Στόχος: Δημιούργησε μια σύντομη, ακριβή περίληψη του παρακάτω αποσπάσματος στα Ελληνικά.\n"
        "Οδηγίες:\n"
        "- Κράτησε βασικά σημεία, συμπεράσματα, αριθμούς/ονόματα (αν υπάρχουν).\n"
        "- Αφαίρεσε επαναλήψεις.\n"
        "- Γράψε 8–12 κουκκίδες (bullets).\n\n"
        "Κείμενο:\n"
        f"{chunk}\n"
    )

def make_final_prompt(bullets_joined: str, target_words: int) -> str:
    return (
        f"{SYSTEM_STYLE}\n\n"
        "Θα σου δώσω συγκεντρωτικές κουκκίδες από πολλά αποσπάσματα ενός εγγράφου.\n"
        f"Στόχος: Γράψε τελική περίληψη ~1 σελίδας στα Ελληνικά (περίπου {target_words} λέξεις).\n"
        "Μορφή:\n"
        "1) Σύντομη εισαγωγή 2–3 προτάσεων.\n"
        "2) 5–8 βασικά σημεία σε κουκκίδες.\n"
        "3) Μια παράγραφος με συμπέρασμα/τελικό μήνυμα.\n"
        "Κανόνες:\n"
        "- Μόνο πληροφορίες που τεκμηριώνονται από τις κουκκίδες.\n"
        "- Καθόλου αναφορές στο ότι είσαι μοντέλο.\n\n"
        "Κουκκίδες:\n"
        f"{bullets_joined}\n"
    )


@dataclass
class LlamaParams:
    n_ctx: int = 8192
    n_threads: int = max(2, os.cpu_count() or 4)
    temperature: float = 0.2
    top_p: float = 0.9
    max_tokens_chunk: int = 400
    max_tokens_final: int = 900


def llm_generate(llm: Llama, prompt: str, max_tokens: int, temperature: float, top_p: float) -> str:
    """
    Uses llama-cpp completion API.
    """
    out = llm(
        prompt,
        max_tokens=max_tokens,
        temperature=temperature,
        top_p=top_p,
        stop=["</s>"],
    )
    return (out["choices"][0]["text"] or "").strip()


# -------------------------
# Main pipeline
# -------------------------
def summarize_docx(
    model_path: str,
    input_docx: str,
    output_docx: str,
    output_txt: str | None,
    target_words: int,
    params: LlamaParams,
) -> None:
    text = read_docx_text(input_docx)
    if not text:
        raise ValueError("Το έγγραφο φαίνεται άδειο ή δεν περιέχει αναγνώσιμο κείμενο.")

    chunks = chunk_text(text, max_chars=9000)
    log.info("Loaded DOCX. Characters=%d, chunks=%d", len(text), len(chunks))

    log.info("Loading model: %s", model_path)
    llm = Llama(
        model_path=model_path,
        n_ctx=params.n_ctx,
        n_threads=params.n_threads,
        verbose=False,
    )

    # 1) Summarize each chunk into bullets
    all_bullets = []
    for i, ch in enumerate(chunks, start=1):
        log.info("Summarizing chunk %d/%d ...", i, len(chunks))
        prompt = make_chunk_prompt(ch)
        bullets = llm_generate(
            llm,
            prompt=prompt,
            max_tokens=params.max_tokens_chunk,
            temperature=params.temperature,
            top_p=params.top_p,
        )
        all_bullets.append(bullets)

    bullets_joined = "\n\n".join(all_bullets)

    # 2) Final 1-page summary from bullets
    log.info("Generating final ~1-page summary ...")
    final_prompt = make_final_prompt(bullets_joined, target_words=target_words)
    summary = llm_generate(
        llm,
        prompt=final_prompt,
        max_tokens=params.max_tokens_final,
        temperature=params.temperature,
        top_p=params.top_p,
    )

    # Save outputs
    title = "Περίληψη (1 σελίδα)"
    write_docx(output_docx, title=title, body=summary)
    log.info("Wrote DOCX: %s", output_docx)

    if output_txt:
        with open(output_txt, "w", encoding="utf-8") as f:
            f.write(summary + "\n")
        log.info("Wrote TXT: %s", output_txt)


def main() -> None:
    parser = argparse.ArgumentParser(description="Summarize Greek DOCX with Meltemi (GGUF) via llama-cpp.")
    parser.add_argument("--model", required=True, help="Path to Meltemi GGUF file (e.g., meltemi-7b-instruct.Q4_K_M.gguf)")
    parser.add_argument("--input", required=True, help="Path to input .docx (Greek)")
    parser.add_argument("--output", required=True, help="Path to output summary .docx")
    parser.add_argument("--output-txt", default=None, help="Optional path to also write summary .txt")
    parser.add_argument("--target-words", type=int, default=520, help="Approx words for 1 page (default ~520)")
    parser.add_argument("--log-level", default=os.getenv("LOG_LEVEL", "INFO"))

    args = parser.parse_args()
    setup_logging(args.log_level)

    params = LlamaParams()
    summarize_docx(
        model_path=args.model,
        input_docx=args.input,
        output_docx=args.output,
        output_txt=args.output_txt,
        target_words=args.target_words,
        params=params,
    )


if __name__ == "__main__":
    main()