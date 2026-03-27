import os
import fitz
from docx import Document

def detect_file_type(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    raise ValueError(f"Unsupported file type: {ext}")

def extract_text_from_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)

def extract_text(file_path: str) -> str:
    file_type = detect_file_type(file_path)
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    if file_type == "docx":
        return extract_text_from_docx(file_path)
    raise ValueError("Unsupported file type")